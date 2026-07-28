"""Regenerate examples/corpus/policy_tables.docx.

A fictional HR/IT policy manual whose answers live inside table cells: it
exists to stress table integrity across chunking strategies. Tables are sized
so that fixed-window chunking demonstrably splits some of them.

Usage: python gen_policy_docx.py
"""

from pathlib import Path

import docx

TARGET = Path(__file__).parent.parent / "corpus" / "policy_tables.docx"


def build() -> docx.Document:
    d = docx.Document()

    def h(level, text):
        d.add_heading(text, level=level)

    def p(text):
        d.add_paragraph(text)

    def table(headers, rows):
        t = d.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        for j, htxt in enumerate(headers):
            t.rows[0].cells[j].text = htxt
        for i, row in enumerate(rows, 1):
            for j, cell in enumerate(row):
                t.rows[i].cells[j].text = str(cell)

    h(1, "Aurora Consulting Group — Global HR and IT Policy Manual")
    p(
        "This manual consolidates the human resources and information technology policies of "
        "Aurora Consulting Group. It applies to all employees, contractors with badge access, and "
        "interns in every country of operation. Where a local law is more favorable to the "
        "employee than a table in this manual, the local law prevails and the difference is "
        "recorded by People Operations in the country annex. The manual is reviewed every "
        "January, and the tables below carry the figures approved for the current fiscal year."
    )

    h(2, "Travel and Expenses")
    p(
        "Business travel is booked through the corporate portal, and the reimbursement caps "
        "depend on the employee's grade and the trip type. Meals outside the caps require a "
        "receipt and a written justification approved by the cost-center owner. The daily "
        "allowance covers meals and incidentals; lodging is booked at cost within the nightly "
        "cap of the destination band. Cities are assigned to bands A, B, or C in the portal, "
        "reviewed twice a year."
    )
    table(
        [
            "Grade",
            "Daily allowance (EUR)",
            "Lodging cap band A (EUR)",
            "Lodging cap band B (EUR)",
            "Flight class short-haul",
            "Flight class long-haul",
        ],
        [
            ["Intern", 45, 140, 105, "Economy", "Economy"],
            ["Analyst", 55, 160, 120, "Economy", "Economy"],
            ["Consultant", 65, 175, 135, "Economy", "Economy"],
            ["Senior Consultant", 75, 190, 150, "Economy", "Premium Economy"],
            ["Manager", 85, 210, 165, "Economy", "Premium Economy"],
            ["Senior Manager", 90, 225, 175, "Economy", "Premium Economy"],
            ["Principal", 95, 240, 185, "Economy", "Business"],
            ["Associate Partner", 105, 270, 200, "Economy", "Business"],
            ["Partner", 120, 300, 220, "Business", "Business"],
        ],
    )
    p(
        "Mileage for personal vehicles is reimbursed at 0.38 euros per kilometer, and rail is "
        "preferred over flights for journeys under four hours door to door. Expense reports are "
        "due within twenty days of trip end; reports filed later than sixty days are rejected "
        "automatically and require a partner exception to reopen."
    )

    h(2, "City Bands for Lodging")
    p(
        "The band of the destination city determines the nightly lodging cap of the travel "
        "table. Cities not listed default to band C, whose cap is eighty-five percent of band B. "
        "The banding follows average business-hotel rates surveyed each April and October."
    )
    table(
        ["City", "Band", "City", "Band"],
        [
            ["Zurich", "A", "Warsaw", "B"],
            ["London", "A", "Prague", "B"],
            ["Paris", "A", "Lisbon", "B"],
            ["New York", "A", "Madrid", "B"],
            ["Milan", "A", "Rome", "B"],
            ["Munich", "A", "Barcelona", "B"],
            ["Amsterdam", "A", "Berlin", "B"],
            ["Copenhagen", "A", "Vienna", "B"],
            ["Stockholm", "A", "Brussels", "B"],
            ["San Francisco", "A", "Frankfurt", "B"],
        ],
    )

    h(2, "Working Hours and On-Call")
    p(
        "Standard contractual hours are set per country, and on-call applies only to roles "
        "designated in the staffing system. On-call compensation combines a standby rate per "
        "rostered day with an intervention rate per activated hour, both differentiated by "
        "weekday and public holiday. Activations are logged automatically by the paging system, "
        "and disputes are resolved against that log."
    )
    table(
        [
            "Roster type",
            "Standby weekday (EUR/day)",
            "Standby holiday (EUR/day)",
            "Intervention rate (EUR/hour)",
            "Max consecutive days",
        ],
        [
            ["Platform on-call", 40, 80, 55, 7],
            ["Client production support", 55, 110, 70, 7],
            ["Security incident response", 70, 140, 90, 5],
            ["Data center hands", 35, 70, 50, 10],
            ["Executive escalation duty", 0, 0, 0, 14],
            ["Facilities emergency", 30, 60, 45, 10],
        ],
    )
    p(
        "An employee may not be rostered for more than the maximum consecutive days shown above, "
        "and a rest day is mandatory after any night with more than four activated hours. "
        "Standby days count toward neither overtime nor the working-time average; activated "
        "hours count toward both."
    )

    h(2, "Leave Entitlements by Country")
    p(
        "Annual leave follows the country of the employment contract, not the country of "
        "assignment. The figures below are working days per calendar year and include any "
        "statutory minimum; they do not include public holidays, which follow the office "
        "calendar of the contract country."
    )
    table(
        [
            "Country",
            "Annual leave (days)",
            "Carry-over limit (days)",
            "Carry-over deadline",
            "Sick pay top-up",
        ],
        [
            ["Italy", 26, 10, "30 June", "100% for 180 days"],
            ["Germany", 28, 5, "31 March", "100% for 42 days"],
            ["France", 27, 6, "31 May", "90% for 90 days"],
            ["Spain", 25, 8, "31 March", "100% for 60 days"],
            ["Poland", 26, 12, "30 September", "80% for 33 days"],
            ["Netherlands", 27, 8, "1 July", "100% for 104 days"],
            ["Sweden", 28, 5, "31 March", "90% for 90 days"],
            ["Denmark", 27, 5, "30 April", "100% for 56 days"],
            ["Austria", 26, 8, "31 March", "100% for 42 days"],
            ["Portugal", 25, 7, "30 April", "100% for 90 days"],
            ["Switzerland", 25, 5, "31 March", "80% for 90 days"],
            ["United Kingdom", 27, 5, "31 March", "100% for 28 days"],
            ["United States", 20, 0, "not applicable", "per state schedule"],
            ["Canada", 21, 5, "31 March", "per province schedule"],
        ],
    )
    p(
        "Unused days beyond the carry-over limit lapse at year end without compensation, except "
        "where local law requires payout. Parental, marriage, and bereavement leave follow the "
        "country annex and are not summarized in this table."
    )

    h(2, "Equipment and Software Provisioning")
    p(
        "Equipment follows the role profile, not the grade. A profile change triggers a "
        "provisioning ticket automatically, and returned devices are wiped following the NIST "
        "800-88 clear standard before reuse or recycling. Peripherals below fifty euros are "
        "stocked in the office supply room and require no ticket."
    )
    table(
        [
            "Role profile",
            "Laptop model class",
            "Refresh cycle (months)",
            "External monitors",
            "Mobile allowance (EUR/month)",
            "Admin rights",
        ],
        [
            ["Standard office", "13-inch ultrabook", 42, 1, 0, "No"],
            ["Consulting delivery", "14-inch performance", 36, 2, 25, "No"],
            ["Data engineering", "16-inch workstation", 30, 2, 25, "On request"],
            ["Design and media", "16-inch workstation", 30, 2, 25, "No"],
            ["Field technician", "Rugged 14-inch", 48, 1, 40, "Yes"],
            ["Quality assurance lab", "15-inch performance", 36, 2, 0, "On request"],
            ["Reception and facilities", "Desktop all-in-one", 60, 1, 0, "No"],
            ["Executive", "14-inch ultralight", 24, 2, 60, "No"],
        ],
    )
    p(
        "Software is licensed per tier: the base tier installs automatically on every device, "
        "the professional tier requires manager approval in the portal, and the restricted tier "
        "requires both a business case and a security review. Personal software installations "
        "are prohibited on managed devices regardless of licence ownership."
    )

    h(2, "Records Retention Schedule")
    p(
        "Records are retained per the schedule below and destroyed in the first quarter "
        "following expiry, under a certificate of destruction for physical media. A legal hold "
        "notified by the general counsel suspends destruction for the records it names until "
        "the hold is released in writing. The schedule lists the controlling copy; convenience "
        "copies are destroyed when their business use ends."
    )
    table(
        ["Record type", "Retention period", "Trigger event", "Storage system"],
        [
            ["Employment contracts", "10 years", "End of employment", "HR vault"],
            ["Payroll records", "10 years", "Fiscal year close", "Payroll archive"],
            ["Recruitment files of non-hired candidates", "18 months", "Position closed", "ATS"],
            ["Working-time records", "5 years", "Calendar year close", "HR vault"],
            ["Expense reports and receipts", "7 years", "Fiscal year close", "Finance archive"],
            ["Client contracts and amendments", "10 years", "Contract end", "Legal repository"],
            ["Client project deliverables", "6 years", "Project closure", "Engagement archive"],
            ["Marketing consent records", "3 years", "Consent withdrawal", "CRM"],
            ["Access logs to production systems", "400 days", "Log creation", "SIEM"],
            ["Email journals", "2 years", "Message date", "Compliance archive"],
            ["Video surveillance footage", "30 days", "Recording", "Facilities NVR"],
            ["Visitor registers", "12 months", "Visit date", "Reception system"],
            ["Health and safety incident reports", "12 years", "Incident date", "HSE system"],
            ["Training completion records", "6 years", "Course completion", "Learning platform"],
        ],
    )
    p(
        "Retention periods satisfy the longest applicable requirement across operating "
        "countries; a country annex may extend but never shorten a period in this table. "
        "Destruction of the controlling copy is logged with date, method, and approver."
    )

    h(2, "Approval Thresholds")
    p(
        "Financial commitments require approval per the thresholds below, computed on the total "
        "contract value including renewals. Splitting a commitment to stay under a threshold is "
        "a disciplinary offence. Emergency purchases follow the same thresholds with approval "
        "collected within two business days after the fact."
    )
    table(
        ["Commitment (EUR)", "Approver", "Second approver", "Procurement involvement"],
        [
            ["Up to 5,000", "Cost-center owner", "None", "No"],
            ["5,001 to 25,000", "Department head", "None", "Advisory"],
            ["25,001 to 100,000", "Department head", "Finance director", "Mandatory"],
            ["100,001 to 500,000", "Finance director", "Managing partner", "Mandatory"],
            ["Above 500,000", "Managing partner", "Partner committee", "Mandatory with tender"],
        ],
    )

    h(2, "Information Classification")
    p(
        "Every document and dataset carries one of four classifications, and the handling rules "
        "follow the classification, not the system it lives in. The author classifies at "
        "creation, and the classification is reviewed at each major revision. When in doubt "
        "between two levels, the higher applies until the data owner rules."
    )
    table(
        [
            "Classification",
            "Examples",
            "Sharing outside Aurora",
            "Encryption at rest",
            "Watermarking",
        ],
        [
            ["Public", "Published brochures", "Free", "Optional", "No"],
            ["Internal", "Process documentation", "With NDA", "Required", "No"],
            ["Confidential", "Client engagement data", "Named recipients only", "Required", "Yes"],
            [
                "Strictly confidential",
                "M&A material, credentials",
                "Prohibited without GC approval",
                "Required with CMK",
                "Yes",
            ],
        ],
    )
    p(
        "Client data inherits at least the Confidential level regardless of the client's own "
        "marking. Reclassification downward requires the data owner and the security office to "
        "agree in writing."
    )

    h(2, "Remote Work Allowances")
    p(
        "Employees on a hybrid or remote contract receive the allowances below, paid monthly "
        "with salary and reviewed annually. The home office setup budget is one-time per "
        "employment relationship and covers furniture and peripherals against receipts."
    )
    table(
        [
            "Contract type",
            "Internet allowance (EUR/month)",
            "Energy allowance (EUR/month)",
            "One-time setup budget (EUR)",
            "Office days required per week",
        ],
        [
            ["Office-based", 0, 0, 0, 4],
            ["Hybrid", 20, 15, 400, 2],
            ["Remote-first", 35, 25, 650, 0],
            ["Client-site", 20, 0, 200, 0],
        ],
    )
    p(
        "Allowances are suspended during unpaid leave months and prorated for part-time "
        "contracts by the contractual percentage. Tax treatment follows the country annex, and "
        "where an allowance is taxable it is grossed up so the net matches this table."
    )

    h(2, "Training Catalog and Certification Support")
    p(
        "Training follows an annual plan agreed at the performance review. Courses in the "
        "catalog below are pre-approved within the stated budget; certifications outside the "
        "catalog need department-head approval. Exam fees are reimbursed on a pass, and one "
        "retake per certification is covered at fifty percent."
    )
    table(
        [
            "Course or certification",
            "Category",
            "Budget cap (EUR)",
            "Study days granted",
            "Validity for refresh",
        ],
        [
            ["Cloud architecture professional", "Technical", 1800, 3, "3 years"],
            ["Data engineering associate", "Technical", 1200, 2, "2 years"],
            ["Security operations analyst", "Technical", 1500, 3, "3 years"],
            ["Agile delivery practitioner", "Methodology", 900, 1, "none"],
            ["Project management professional", "Methodology", 1600, 3, "3 years"],
            ["Negotiation fundamentals", "Consulting skills", 700, 1, "none"],
            ["Executive presence workshop", "Consulting skills", 1100, 1, "none"],
            ["Language course, per level", "Language", 600, 0, "none"],
            ["First aid and safety officer", "Compliance", 350, 1, "2 years"],
            ["Data protection specialist", "Compliance", 950, 2, "3 years"],
        ],
    )
    p(
        "Study days are working days released from staffing, capped at five per calendar year "
        "across all courses. Unused training budget does not carry over, and the learning "
        "platform records completions in the retention schedule of this manual."
    )

    h(2, "Incident Severity and Response")
    p(
        "Security and IT incidents are classified at intake by the on-duty responder, and the "
        "response clock starts at classification. Severity drives both the response time and who "
        "must be informed; the communication rules are mandatory even when the technical "
        "response is trivial."
    )
    table(
        [
            "Severity",
            "Definition",
            "Response start",
            "Status update cadence",
            "Mandatory notification",
        ],
        [
            [
                "S1",
                "Confirmed data breach or client production down",
                "15 minutes",
                "Every hour",
                "CISO, GC, affected client partner",
            ],
            ["S2", "Internal production down, no data exposure", "1 hour", "Every 4 hours", "CISO"],
            [
                "S3",
                "Degraded service with workaround",
                "4 business hours",
                "Daily",
                "Service owner",
            ],
            ["S4", "Cosmetic or single-user issue", "2 business days", "On resolution", "None"],
        ],
    )
    p(
        "An S1 involving personal data additionally triggers the privacy procedure, whose "
        "regulatory notification assessment must conclude within forty-eight hours of "
        "classification. Post-incident reviews are mandatory for S1 and S2 and produce actions "
        "tracked to closure in the risk register."
    )

    return d


if __name__ == "__main__":
    build().save(TARGET)
    print(f"wrote {TARGET}")
