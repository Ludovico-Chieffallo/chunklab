"""Document shapes that real files have and hand-written fixtures do not.

Found by running the loaders over a real 10-K filing and a real annual report.
Both broke `structure` chunking *silently* — it kept producing chunks, they were
just packed to `max_tokens` with no regard for the document's sections.

The fixtures here are synthetic reproductions of those shapes; the source files
are private and never enter the repository.
"""

import docx
import pytest

from chunklab.chunkers.structure import StructureChunker
from chunklab.config import default_config
from chunklab.loaders.docx import DocxLoader
from chunklab.models import Document, Element, Question
from chunklab.runner import run_evaluation
from chunklab.text_utils import count_tokens

BODY = (
    "Revenue increased across every reportable segment during the fiscal year, "
    "driven by higher unit volumes and favourable pricing in the second half. "
    "Operating expenses grew more slowly than revenue in the same period. "
)


def _write_docx(path, blocks):
    """blocks: (text, bold, style_name) triples."""
    document = docx.Document()
    for text, bold, style in blocks:
        paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
    document.save(str(path))
    return path


# --- headings that carry no heading style --------------------------------------


def test_bold_short_paragraphs_are_recognised_as_headings(tmp_path):
    """A real annual report had 961 paragraphs, every one styled `Normal (Web)`,
    while being visually structured with 171 bold headings."""
    path = _write_docx(
        tmp_path / "report.docx",
        [
            ("OUR CULTURE", True, None),
            (BODY, False, None),
            ("Share Repurchases", True, None),
            (BODY, False, None),
        ],
    )

    document = DocxLoader().load(path)

    headings = [e for e in document.elements if e.type == "heading"]
    assert [e.text for e in headings] == ["## OUR CULTURE", "### Share Repurchases"]
    assert document.metadata["inferred_headings"] == 2


def test_all_caps_heading_outranks_title_case(tmp_path):
    path = _write_docx(
        tmp_path / "r.docx", [("BUSINESS", True, None), ("Embracing Our Future", True, None)]
    )

    levels = [e.level for e in DocxLoader().load(path).elements if e.type == "heading"]

    assert levels == [2, 3]


@pytest.mark.parametrize(
    ("text", "bold", "why"),
    [
        (BODY, True, "a long bold paragraph is emphasis, not a heading"),
        ("We grew revenue this year.", True, "a bold sentence ends in punctuation"),
        ("OUR CULTURE", False, "not bold at all"),
    ],
)
def test_things_that_are_not_headings(tmp_path, text, bold, why):
    path = _write_docx(tmp_path / "r.docx", [(text, bold, None)])

    headings = [e for e in DocxLoader().load(path).elements if e.type == "heading"]

    assert headings == [], why


def test_partially_bold_paragraph_is_not_a_heading(tmp_path):
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Net sales ").bold = True
    paragraph.add_run("were higher")
    path = tmp_path / "r.docx"
    document.save(str(path))

    assert [e for e in DocxLoader().load(path).elements if e.type == "heading"] == []


def test_real_heading_styles_still_win(tmp_path):
    path = _write_docx(tmp_path / "r.docx", [("Part I", False, "Heading 1"), (BODY, False, None)])

    document = DocxLoader().load(path)

    headings = [e for e in document.elements if e.type == "heading"]
    assert [(e.level, e.text) for e in headings] == [(1, "# Part I")]
    assert document.metadata["inferred_headings"] == 0


# --- outlines that a converter shifted deeper ----------------------------------


def _doc_with_heading_levels(levels: list[int]) -> Document:
    parts: list[str] = []
    elements: list[Element] = []
    offset = 0
    for index, level in enumerate(levels):
        heading = "#" * level + f" Section {index}"
        parts.append(heading + "\n\n")
        elements.append(
            Element(
                type="heading",
                text=heading,
                char_span=(offset, offset + len(heading)),
                level=level,
            )
        )
        offset += len(heading) + 2
        parts.append(BODY + "\n\n")
        offset += len(BODY) + 2
    return Document(
        id="doc", source_path="doc.md", text="".join(parts), elements=elements, metadata={}
    )


def test_split_depth_follows_where_the_headings_actually_are():
    """Regression: pymupdf4llm gave level 5 to 261 of 268 headings in a real 10-K,
    so a fixed `<= 3` saw three headings in a 66k-token filing."""
    assert StructureChunker._split_depth(_doc_with_heading_levels([5] * 20).elements) == 5


def test_split_depth_is_unchanged_for_a_conventional_outline():
    elements = _doc_with_heading_levels([1, 2, 2, 3, 2, 3]).elements
    assert StructureChunker._split_depth(elements) == 3


def test_deep_uniform_headings_still_produce_sections():
    document = _doc_with_heading_levels([5] * 12)

    chunks = StructureChunker(max_tokens=800).chunk(document)

    assert len(chunks) == 12, "a shifted outline collapsed into one packed chunk"
    assert max(count_tokens(c.text) for c in chunks) < 200


# --- telling the user when structure has nothing to work with ------------------


def test_structure_without_headings_is_reported():
    config = default_config()
    config.embedding.backend = "fake"
    documents = [
        Document(id="flat", source_path="flat.txt", text=BODY * 6, elements=[], metadata={})
    ]
    questions = [
        Question(
            id="q1",
            query="How did revenue move?",
            gold_snippets=["Revenue increased across every reportable segment"],
        )
    ]

    report = run_evaluation(documents, questions, config)

    assert any("no detectable headings" in w and "flat" in w for w in report.warnings)


def test_no_such_warning_when_headings_exist():
    config = default_config()
    config.embedding.backend = "fake"
    documents = [_doc_with_heading_levels([2, 2, 3])]
    questions = [
        Question(
            id="q1",
            query="How did revenue move?",
            gold_snippets=["Revenue increased across every reportable segment"],
        )
    ]

    report = run_evaluation(documents, questions, config)

    assert not any("no detectable headings" in w for w in report.warnings)
